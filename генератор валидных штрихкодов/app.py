from PyQt4 import QtGui
from PyQt4.QtGui import *
from PyQt4 import QtCore
import sys
from PIL import Image
from PIL import ImageDraw
from PIL import ImageFont
import os
from docx import Document
from docx.shared import Inches


FONT_PATH = "arialbd.ttf"
LINE_WIDTH = 2
WIDTH = 240
HEIGHT = 110
SIDE_LINE_WIDTH = 97
MAIN_LINE_WIDTH = 87
FONT_SIZE = 18


CODE_TYPES = ['EAN-13', 'EAN-8', 'EAN-13+2', 'EAN-13+5', 'ISBN']

group_A = {
    '0': '0001101',
    '1': '0011001',
    '2': '0010011',
    '3': '0111101',
    '4': '0100011',
    '5': '0110001',
    '6': '0101111',
    '7': '0111011',
    '8': '0110111',
    '9': '0001011'
    }
group_B = {
    '0': '0100111',
    '1': '0110011',
    '2': '0011011',
    '3': '0100001',
    '4': '0011101',
    '5': '0111001',
    '6': '0000101',
    '7': '0010001',
    '8': '0001001',
    '9': '0010111'
    }
group_C = {
    '0': '1110010',
    '1': '1100110',
    '2': '1101100',
    '3': '1000010',
    '4': '1011100',
    '5': '1001110',
    '6': '1010000',
    '7': '1000100',
    '8': '1001000',
    '9': '1110100'
    }
combination = {
    '0':'AAAAAA',
    '1':'AABABB',
    '2':'AABBAB',
    '3':'AABBBA',
    '4':'ABAABB',
    '5':'ABBAAB',
    '6':'ABBBAA',
    '7':'ABABAB',
    '8':'AAABBA',
    '9':'ABBABA'
    }


class Generator(QtGui.QWidget):
    idEdit = None
    contryEdit = None
    countEdit = None
    priceEdit = None
    emailEdit = None
    log = None
    image = None
    
    progress = None
    
    def __init__(self):
        super(Generator, self).__init__()
        
        self.initUI()
        
    def initUI(self):
        
        codeS_title = QtGui.QLabel('Type')
        title = QtGui.QLabel('Country')
        author = QtGui.QLabel('ID')
        review = QtGui.QLabel('Count')
        price = QtGui.QLabel('Price')
        email = QtGui.QLabel('Email')
        
        self.contryEdit = QtGui.QLineEdit()
        self.idEdit = QtGui.QLineEdit()
        self.log = QtGui.QTextEdit()
        self.countEdit = QtGui.QLineEdit()
        self.priceEdit = QtGui.QLineEdit()
        self.emailEdit = QtGui.QLineEdit()
        
        self.image = QtGui.QLabel("window")
        self.image.setGeometry(10, 10, WIDTH, HEIGHT)
        #use full ABSOLUTE path to the image, not relative
        self.image.setPixmap(QtGui.QPixmap( "barcode_scanner_icon.png"))

        button = QtGui.QPushButton()
        button.setText("GENERATE")
        button.clicked.connect(self.handleButton)
        
        toHtml = QtGui.QPushButton()
        toHtml.setText("HTML")
        toHtml.clicked.connect(self.toHTML)
        
        
          
        self.progress = QtGui.QProgressBar();
        self.progress.setMaximum(0)
        self.progress.setMinimum(0)
        self.progress.setValue(0)
        
        listWidget = QComboBox()
        listWidget.addItems(CODE_TYPES)
            
        grid = QtGui.QGridLayout()
        grid.setSpacing(20)

        grid.addWidget(self.image, 1, 1)
        
        #grid.addWidget(codeS_title, 1, 0)
        #grid.addWidget(listWidget, 1, 1)
        
        grid.addWidget(title, 2, 0)
        grid.addWidget(self.contryEdit, 2, 1)

        grid.addWidget(author, 3, 0)
        grid.addWidget(self.idEdit, 3, 1)

        grid.addWidget(review, 4, 0)
        grid.addWidget(self.countEdit, 4, 1)

        grid.addWidget(price, 5, 0)
        grid.addWidget(self.priceEdit, 5, 1)
        
        grid.addWidget(email, 6, 0)
        grid.addWidget(self.emailEdit, 6, 1)
        
        grid.addWidget(button, 7, 1)
        
        grid.addWidget(self.progress, 8, 1)
        
        grid.addWidget(email, 9, 0)
        grid.addWidget(self.log, 9, 1,5,1)
        
        grid.addWidget(toHtml, 10, 0)
        
        self.setLayout(grid) 
        
        self.setGeometry(300, 300, 350, 100)
        self.setWindowTitle('BARCODE GENERATOR')    
        self.show()
    def toHTML(self):  
        f = open("file_name.html" , "w")
        f.write(self.log.toHtml())
        f.close()
        os.startfile("file_name.html")
        
    def handleButton(self):        
        country = self.contryEdit.text()         
        codeBusiness = self.idEdit.text()
        N = int(self.countEdit.text()) 
        
        self.progress.setMaximum(N)      
        
        if not os.path.isdir("codes"):
            try:
                os.makedirs('codes')
            except OSError:
                pass
                print "Error creating direcory [codes]"
        
        index = 1
        COL_SIZE = 4
        document = Document()
        tbl = document.add_table(rows=1, cols=4)
        x = 0
        while index<=N:
            codeProduct = self.recAdd(str(index))#5
            self.draw(str(country+codeBusiness+codeProduct+str(self.checksum(country+codeBusiness+codeProduct))),"codes\\"+str(index)+".jpg")
            self.progress.setValue(index)
            if x%COL_SIZE==0:
                row_cells = tbl.add_row().cells
                x = 0

            paragraph = row_cells[x].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("codes\\"+str(index)+".jpg",width=Inches(1.40))
            index+=1
            x+=1
            
        self.image.setPixmap(QtGui.QPixmap( "codes\\1.jpg"))
        document.save("demo.docx")
        
    def closeEvent(self, event):
        print self.idEdit.text()
        
    def recSearch(self,t):
        t+=1
        if t%10==0:
            return t
        else:        
            t = self.recSearch(t)
        return t

    def recAdd(self,x):
        if (len(x)<5):
            x="0"+x
            x = self.recAdd(x)
        return x        

    def checksum(self,x):
        i = 0
        sum_1 = 0
        sum_2 = 0
        while i<len(x):
            sum_1+=int(x[i])
            i+=2
        i=1
        while i<len(x):
            sum_2+=int(x[i])
            i+=2
        sum_2 *=3
        sum_1 +=sum_2
        return self.recSearch(sum_1)-sum_1

    def format(self,x,n):
        if (len(x)<int(n)):
            x = "0"+x
            x = self.format(self,x,n)
        return x
    
    def draw(self,code,path_f):
        image = Image.new("RGB", (WIDTH, HEIGHT), (255, 255, 255))
        draw = ImageDraw.Draw(image)  

        comb = code[0]
        part_1 = code[1:7:1]
        part_2 = code[7:13:1]
        path = combination[str(int(comb))]
        index = 0
        step = 22
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
        step += 2*LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
        step += LINE_WIDTH       
        self.log.append(comb+"["+part_1+"]["+part_2+"]["+path+"CCCCCC] - successed")
        for i in path:

            if i == 'A': 
                for k in group_A[str(int(part_1[index]))]:
                    if (k == '1'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
                    if (k == '0'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255),width=LINE_WIDTH)
                    step += LINE_WIDTH            
            if i == 'B':
                for k in group_B[str(int(part_1[index]))]:
                    if (k == '1'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
                    if (k == '0'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255),width=LINE_WIDTH)
                    step += LINE_WIDTH
            index += 1
        step += LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
        step += 2*LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH) 
        step += 2*LINE_WIDTH

        index = 0
        while index < 6:
            for k in group_C[str(int(part_2[index]))]:
                if (k == '1'):
                    draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
                if (k == '0'):
                    draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255),width=LINE_WIDTH)
                step += LINE_WIDTH
            index += 1
        #step += LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)
        step += 2*LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0),width=LINE_WIDTH)

        font = ImageFont.truetype(FONT_PATH, FONT_SIZE)
        draw.text((6, MAIN_LINE_WIDTH+2), comb, (1, 1, 1), font=font)
        k = 45
        for i in range(len(part_1)):
             draw.text((k, MAIN_LINE_WIDTH+2), part_1[i:i+1:1], (1, 1, 1), font=font)
             k+=10

        k = 130
        for i in range(len(part_2)):
             draw.text((k, MAIN_LINE_WIDTH+2), part_2[i:i+1:1], (1, 1, 1), font=font)
             k+=10
        percentage_ = 100
        size_ = (image.size[0] * percentage_ / 100, image.size[1] * percentage_ / 100)
        image=image.resize(size_, Image.ANTIALIAS)
        image.save(path_f, "JPEG")
        #image.show()   
        del draw
        return 0
        
def main():
    
    app = QtGui.QApplication(sys.argv)
    ex = Generator()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main() 