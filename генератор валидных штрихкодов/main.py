from PIL import Image
from PIL import ImageDraw
from PIL import ImageFont
from PyQt4 import QtCore
from PyQt4 import QtGui
from PyQt4.QtGui import *
from docx import Document
from docx.shared import Inches
from email import Encoders
from email.MIMEBase import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os
import smtplib
import sys
import datetime
from configparser import ConfigParser

FONT_PATH = "arialbd.ttf"
LINE_WIDTH = 2
WIDTH = 240
HEIGHT = 110
SIDE_LINE_WIDTH = 97
MAIN_LINE_WIDTH = 87
FONT_SIZE = 18


CODE_TYPES = ['EAN-13', 'EAN-8', 'EAN-13+2', 'EAN-13+5', 'ISBN']

group_A = {
    '0': '0001101',
    '1': '0011001',
    '2': '0010011',
    '3': '0111101',
    '4': '0100011',
    '5': '0110001',
    '6': '0101111',
    '7': '0111011',
    '8': '0110111',
    '9': '0001011'
    }
group_B = {
    '0': '0100111',
    '1': '0110011',
    '2': '0011011',
    '3': '0100001',
    '4': '0011101',
    '5': '0111001',
    '6': '0000101',
    '7': '0010001',
    '8': '0001001',
    '9': '0010111'
    }
group_C = {
    '0': '1110010',
    '1': '1100110',
    '2': '1101100',
    '3': '1000010',
    '4': '1011100',
    '5': '1001110',
    '6': '1010000',
    '7': '1000100',
    '8': '1001000',
    '9': '1110100'
    }
combination = {
    '0':'AAAAAA',
    '1':'AABABB',
    '2':'AABBAB',
    '3':'AABBBA',
    '4':'ABAABB',
    '5':'ABBAAB',
    '6':'ABBBAA',
    '7':'ABABAB',
    '8':'AAABBA',
    '9':'ABBABA'
    }


class Generator(QtGui.QWidget): 
    def __init__(self):
        super(Generator, self).__init__()
        
        self.initUI()
        
    def initUI(self):  
        param = self.read_config()
        self.codeS_title = QtGui.QLabel('Type')
        self.title = QtGui.QLabel('Country')
        self.author = QtGui.QLabel('ID')
        self.review = QtGui.QLabel('Count')
        self.price = QtGui.QLabel('Price')
        self.email = QtGui.QLabel('Email')        
        self.contryEdit = QtGui.QLineEdit()
        self.contryEdit.setText(str(param['country']))
        self.idEdit = QtGui.QLineEdit()
        self.idEdit.setText(str(param['id']))
        self.log = QtGui.QTextEdit()
        self.countEdit = QtGui.QLineEdit()
        self.countEdit.setText(str(param['count']))
        self.priceEdit = QtGui.QLineEdit()
        self.priceEdit.setText(str(param['price']))
        self.emailEdit = QtGui.QLineEdit()     
        self.emailEdit.setText(str(param['email']))
        self.image = QtGui.QLabel("window")
        self.image.setGeometry(10, 10, WIDTH, HEIGHT)
        self.image.setPixmap(QtGui.QPixmap("barcode_scanner_icon.png"))
        self.button = QtGui.QPushButton()
        self.button.setText("GENERATE")
        self.button.clicked.connect(self.handleButton)
        self.toHtml = QtGui.QPushButton()
        self.toHtml.setText("SEND")
        self.toHtml.clicked.connect(self.toHTML)
        self.clear = QtGui.QPushButton()
        self.clear.setText("HELP")
        self.clear.clicked.connect(self.HELPMe)
        self.watermark = QtGui.QPushButton()
        self.watermark.setText("Watermark")
        self.watermark.clicked.connect(self.addWaterMark)
        self.progress = QtGui.QProgressBar();
        self.progress.setMaximum(0)
        self.progress.setMinimum(0)
        self.progress.setValue(0)     
        self.listWidget = QComboBox()
        self.listWidget.addItems(CODE_TYPES) 
        self.grid = QtGui.QGridLayout(self)
        self.grid.setSpacing(20) 
        self.listWidget = QComboBox()
        self.listWidget.addItems(CODE_TYPES)
        self.grid.addWidget(self.image, 1, 1)
        self.grid.addWidget(self.codeS_title, 2, 0)
        self.grid.addWidget(self.listWidget, 2, 1)
        self.grid.addWidget(self.title, 3, 0)
        self.grid.addWidget(self.contryEdit, 3, 1)
        self.grid.addWidget(self.author, 4, 0)
        self.grid.addWidget(self.idEdit, 4, 1)
        self.grid.addWidget(self.review, 5, 0)
        self.grid.addWidget(self.countEdit, 5, 1)
        self.grid.addWidget(self.price, 6, 0)
        self.grid.addWidget(self.priceEdit, 6, 1)
        self.grid.addWidget(self.email, 7, 0)
        self.grid.addWidget(self.emailEdit, 7, 1)
        self.grid.addWidget(self.button, 8, 1)
        self.grid.addWidget(self.progress, 9, 1)
        self.grid.addWidget(self.log, 10, 1, 5, 1)
        self.grid.addWidget(self.toHtml, 11, 0)
        self.grid.addWidget(self.watermark, 12, 0)
        self.grid.addWidget(self.clear, 13, 0)
        self.setLayout(self.grid) 
        self.setGeometry(300, 300, 450, 100)
        self.setWindowTitle('BARCODE GENERATOR')    
        self.show()
        
    def HELPMe(self):
        try:
             os.startfile("help.html")
        except Exception as e:
            print e
           
        
    def addWaterMark(self):        
        fname = QtGui.QFileDialog.getOpenFileName(self, 'Open watermark Image', '', "Image Files (*.png)")
        if (not os.path.exists(fname)):
            print "File not exist"
            return 0        
        im2 = Image.open(str(fname))
        if not os.path.isdir("waterMark"):
            try:
                os.makedirs('waterMark')
            except OSError:
                print "Error creating direcory [waterMark]"
        for i in  os.listdir('waterMark'):
            print i
            if (i.find(".jpg") != -1 or i.find(".png") != -1):
                im1 = Image.open(str("waterMark\\" + i))
                im1.paste(im2, (im1.size[0] / 2-im2.size[0] / 2, im1.size[1] / 2-im2.size[1] / 2), im2) 
                im1.save(str("waterMark\\" + i))
        
    def read_config(self,filename='config.ini', section='program'):
        try:
            parser = ConfigParser()           
            parser.read(filename)           

            # get section, default to mysql
            params = {}
            if parser.has_section(section):
                items = parser.items(section)
                for item in items:
                    params[item[0]] = item[1]
            else:
                raise Exception('{0} not found in the {1} file'.format(section, filename))
            return params
        except Exception as e:
            print e        
    def toHTML(self):
        date = str(datetime.datetime.time(datetime.datetime.now())).replace(":","_").replace(".","_")
        f = open("log_info_"+date+".html", "w")
        f.write(self.log.toHtml())
        f.close()
        param = self.read_config()
        os.startfile("log_info_"+date+".html")
        me = param['my_mail']
        you =  str(self.emailEdit.text())

        text = param['text']
        subj = 'latest barcodes '+date

        server = "smtp.gmail.com"
        port = 25
        user_name = str(param['user_name'])
        user_passwd = str(param['user_passwd'])

        msg = MIMEMultipart('mixed')
        msg['Subject'] = str(subj)
        msg['From'] = str(me)
        msg['To'] = str(you)

        for i in  os.listdir('docs'):
            basename = "docs\\"+i
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(basename, "rb").read())
            Encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment; filename="%s"' % basename)
            msg.attach(part)

        try:
            s = smtplib.SMTP(server, port)
            s.ehlo()
            s.starttls()
            s.ehlo()
            s.login(user_name, user_passwd)
            s.sendmail(me, you, msg.as_string())
            s.quit()
        except Exception as e:
            print e
            os.startfile("https://www.google.com/settings/security/lesssecureapps")
            
        
    def handleButton(self):    
        self.log.clear()
        country = self.format(self.contryEdit.text(),3)        
        codeBusiness = self.format(self.idEdit.text(),4)                    
        N = int(self.countEdit.text()) 
        
        self.progress.setMaximum(N)      
        
        if not os.path.isdir("codes"):
            try:
                os.makedirs('codes')
            except OSError:
                pass
                print "Error creating direcory [codes]"
        
        index = 1
        COL_SIZE = 4
        document = Document()
        tbl = document.add_table(rows=1, cols=4)
        x = 0
        while index <= N:
            codeProduct = self.recAdd(str(index))#5
            self.draw(str(country + codeBusiness + codeProduct + str(self.checksum(country + codeBusiness + codeProduct))), "codes\\" + str(index) + ".jpg")
            self.progress.setValue(index)
            if x % COL_SIZE == 0:
                row_cells = tbl.add_row().cells
                x = 0

            paragraph = row_cells[x].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("codes\\" + str(index) + ".jpg", width=Inches(1.40))
            index += 1
            x += 1
            
        self.image.setPixmap(QtGui.QPixmap("codes\\1.jpg"))
        date = str(datetime.datetime.time(datetime.datetime.now())).replace(":","_").replace(".","_")
        
        if not os.path.isdir("docs"):
            try:
                os.makedirs('docs')
            except OSError:
                print "Error creating direcory [docs]"
        document.save("docs\\barcodes"+date+".docx")
        os.startfile("docs\\barcodes"+date+".docx")
        
    def closeEvent(self, event):
        print self.idEdit.text()
        
    def recSearch(self, t):
        t += 1
        if t % 10 == 0:
            return t
        else:        
            t = self.recSearch(t)
        return t

    def recAdd(self, x):
        if (len(x) < 5):
            x = "0" + x
            x = self.recAdd(x)
        return x        

    def checksum(self, x):
        i = 0
        sum_1 = 0
        sum_2 = 0
        while i < len(x):
            sum_1 += int(x[i])
            i += 2
        i = 1
        while i < len(x):
            sum_2 += int(x[i])
            i += 2
        sum_2 *= 3
        sum_1 += sum_2
        return self.recSearch(sum_1)-sum_1

    def format(self, x, n):
        if (len(x) < int(n)):
            x = "0" + x
            x = self.format(self, x, n)
        return x
    
    def draw(self, code, path_f):
        image = Image.new("RGB", (WIDTH, HEIGHT), (255, 255, 255))
        draw = ImageDraw.Draw(image)  
        
        comb = code[0]
        part_1 = code[1:7:1]
        part_2 = code[7:13:1]
        path = combination[str(int(comb))]
        index = 0
        step = 22
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
        step += 2 * LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
        step += LINE_WIDTH       
        self.log.append(comb + "[" + part_1 + "][" + part_2 + "][" + path + "CCCCCC] - successed")
        for i in path:

            if i == 'A': 
                for k in group_A[str(int(part_1[index]))]:
                    if (k == '1'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
                    if (k == '0'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255), width=LINE_WIDTH)
                    step += LINE_WIDTH            
            if i == 'B':
                for k in group_B[str(int(part_1[index]))]:
                    if (k == '1'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
                    if (k == '0'):
                        draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255), width=LINE_WIDTH)
                    step += LINE_WIDTH
            index += 1
        step += LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
        step += 2 * LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH) 
        step += 2 * LINE_WIDTH

        index = 0
        while index < 6:
            for k in group_C[str(int(part_2[index]))]:
                if (k == '1'):
                    draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
                if (k == '0'):
                    draw.line((step, 0, step, MAIN_LINE_WIDTH), fill=(255, 255, 255), width=LINE_WIDTH)
                step += LINE_WIDTH
            index += 1
        #step += LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)
        step += 2 * LINE_WIDTH
        draw.line((step, 0, step, SIDE_LINE_WIDTH), fill=(0, 0, 0), width=LINE_WIDTH)

        font = ImageFont.truetype(FONT_PATH, FONT_SIZE)
        draw.text((6, MAIN_LINE_WIDTH + 2), comb, (1, 1, 1), font=font)
        k = 45
        for i in range(len(part_1)):
            draw.text((k, MAIN_LINE_WIDTH + 2), part_1[i:i + 1:1], (1, 1, 1), font=font)
            k += 10

        k = 130
        for i in range(len(part_2)):
            draw.text((k, MAIN_LINE_WIDTH + 2), part_2[i:i + 1:1], (1, 1, 1), font=font)
            k += 10
        percentage_ = 100
        size_ = (image.size[0] * percentage_ / 100, image.size[1] * percentage_ / 100)
        image = image.resize(size_, Image.ANTIALIAS)
        image.save(path_f, "JPEG")
        #image.show()   
        del draw
        return 0
        
def main():
    
    app = QtGui.QApplication(sys.argv)
    ex = Generator()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main() 